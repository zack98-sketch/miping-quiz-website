import io
from datetime import datetime, timedelta
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func
from app.database import get_db
from app.models import (
    Question, QuizSession, AnswerRecord, ErrorBookItem, User, SessionStatus
)
from app.dependencies import get_current_user
from typing import Optional

router = APIRouter(prefix="/api/analysis", tags=["学习分析"])


@router.get("/mastery")
def get_mastery(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Get all answer records for this user
    records = db.query(AnswerRecord).join(QuizSession).filter(QuizSession.user_id == current_user.id).all()
    
    # Group by knowledge point
    kp_stats = {}
    for r in records:
        q = db.query(Question).filter(Question.id == r.question_id).first()
        if not q:
            continue
        kp = q.knowledge_point
        if kp not in kp_stats:
            kp_stats[kp] = {"total": 0, "correct": 0}
        kp_stats[kp]["total"] += 1
        if r.is_correct == 1:
            kp_stats[kp]["correct"] += 1
    
    result = []
    for kp, stats in kp_stats.items():
        rate = (stats["correct"] / stats["total"] * 100) if stats["total"] > 0 else 0
        result.append({
            "knowledge_point": kp,
            "total_count": stats["total"],
            "correct_count": stats["correct"],
            "error_count": stats["total"] - stats["correct"],
            "mastery_rate": round(rate, 1),
        })
    
    result.sort(key=lambda x: x["mastery_rate"])
    return {"items": result}


@router.get("/weak-points")
def get_weak_points(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    mastery = get_mastery(current_user, db)
    weak = [item for item in mastery["items"] if item["mastery_rate"] < 60]
    return {"items": weak}


@router.get("/recommendations")
def get_recommendations(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Get error book knowledge points
    error_kps = db.query(Question.knowledge_point).join(
        ErrorBookItem, ErrorBookItem.question_id == Question.id
    ).filter(ErrorBookItem.user_id == current_user.id).distinct().all()
    
    recommendations = []
    for (kp,) in error_kps:
        if not kp:
            continue
        # Count unanswered questions in this knowledge point
        answered_ids = db.query(AnswerRecord.question_id).join(QuizSession).filter(
            QuizSession.user_id == current_user.id,
            AnswerRecord.is_correct == 1,
        ).subquery()
        
        available = db.query(Question).filter(
            Question.knowledge_point == kp,
            Question.is_active == True,
            Question.id.notin_(db.query(AnswerRecord.question_id).join(QuizSession).filter(
                QuizSession.user_id == current_user.id,
                AnswerRecord.is_correct == 1,
            )),
        ).count()
        
        recommendations.append({
            "knowledge_point": kp,
            "reason": f"该知识点存在错题，建议加强练习",
            "question_count": available,
        })
    
    return {"items": recommendations}


@router.get("/progress")
def get_progress(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    total_questions = db.query(AnswerRecord).join(QuizSession).filter(QuizSession.user_id == current_user.id).count()
    total_correct = db.query(AnswerRecord).join(QuizSession).filter(QuizSession.user_id == current_user.id, AnswerRecord.is_correct == 1).count()
    total_errors = total_questions - total_correct
    accuracy = (total_correct / total_questions * 100) if total_questions > 0 else 0
    
    # Study days
    first_record = db.query(QuizSession).filter(QuizSession.user_id == current_user.id).order_by(QuizSession.started_at).first()
    study_days = 0
    if first_record and first_record.started_at:
        study_days = (datetime.utcnow() - first_record.started_at).days + 1
    
    total_sessions = db.query(QuizSession).filter(QuizSession.user_id == current_user.id, QuizSession.status == SessionStatus.completed).count()
    
    return {
        "total_questions": total_questions,
        "total_correct": total_correct,
        "total_errors": total_errors,
        "overall_accuracy": round(accuracy, 1),
        "study_days": study_days,
        "total_sessions": total_sessions,
    }


@router.get("/history")
def get_history(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(QuizSession).filter(QuizSession.user_id == current_user.id, QuizSession.status == SessionStatus.completed)
    if start_date:
        query = query.filter(QuizSession.started_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(QuizSession.started_at <= datetime.fromisoformat(end_date))
    
    total = query.count()
    sessions = query.order_by(QuizSession.started_at.desc()).offset((page - 1) * page_size).limit(page_size).all()
    
    items = []
    for s in sessions:
        accuracy = (s.correct_count / s.total_count * 100) if s.total_count > 0 else 0
        items.append({
            "session_id": s.id,
            "mode": s.mode.value,
            "total_count": s.total_count,
            "correct_count": s.correct_count,
            "accuracy": round(accuracy, 1),
            "time_spent": s.time_spent,
            "started_at": s.started_at.isoformat() if s.started_at else None,
            "completed_at": s.completed_at.isoformat() if s.completed_at else None,
        })
    
    return {"total": total, "page": page, "page_size": page_size, "items": items}


@router.get("/trend")
def get_trend(days: int = Query(30, ge=1, le=365), current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    since = datetime.utcnow() - timedelta(days=days)
    sessions = db.query(QuizSession).filter(
        QuizSession.user_id == current_user.id,
        QuizSession.status == SessionStatus.completed,
        QuizSession.started_at >= since,
    ).order_by(QuizSession.started_at).all()
    
    # Group by date
    daily = {}
    for s in sessions:
        date_key = s.started_at.strftime("%Y-%m-%d") if s.started_at else "unknown"
        if date_key not in daily:
            daily[date_key] = {"correct": 0, "total": 0}
        daily[date_key]["correct"] += s.correct_count
        daily[date_key]["total"] += s.total_count
    
    result = []
    for date, stats in sorted(daily.items()):
        accuracy = (stats["correct"] / stats["total"] * 100) if stats["total"] > 0 else 0
        result.append({"date": date, "accuracy": round(accuracy, 1), "count": stats["total"]})
    
    return {"items": result}


@router.get("/export/pdf")
def export_pdf(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        
        # Simple PDF generation
        c.setFont("Helvetica", 12)
        c.drawString(100, 800, "Answer Records Export")
        
        sessions = db.query(QuizSession).filter(
            QuizSession.user_id == current_user.id,
            QuizSession.status == SessionStatus.completed,
        ).order_by(QuizSession.started_at.desc()).limit(50).all()
        
        y = 760
        for s in sessions:
            if y < 50:
                c.showPage()
                y = 800
            accuracy = (s.correct_count / s.total_count * 100) if s.total_count > 0 else 0
            text = f"Session {s.id}: {s.correct_count}/{s.total_count} ({accuracy:.1f}%) - {s.mode.value}"
            c.drawString(50, y, text)
            y -= 20
        
        c.save()
        buf.seek(0)
        
        return StreamingResponse(buf, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=answer_records.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF导出失败: {str(e)}")


@router.get("/export/word")
def export_word(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading("答题记录导出", 0)
        
        sessions = db.query(QuizSession).filter(
            QuizSession.user_id == current_user.id,
            QuizSession.status == SessionStatus.completed,
        ).order_by(QuizSession.started_at.desc()).limit(50).all()
        
        for s in sessions:
            accuracy = (s.correct_count / s.total_count * 100) if s.total_count > 0 else 0
            doc.add_paragraph(f"答题会话 {s.id}: {s.correct_count}/{s.total_count} ({accuracy:.1f}%) - 模式: {s.mode.value}")
            
            records = db.query(AnswerRecord).filter(AnswerRecord.session_id == s.id).all()
            for r in records:
                q = db.query(Question).filter(Question.id == r.question_id).first()
                if q:
                    status = "正确" if r.is_correct == 1 else "错误"
                    doc.add_paragraph(f"  [{status}] {q.content[:80]}...", style="List Bullet")
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=answer_records.docx"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word导出失败: {str(e)}")
